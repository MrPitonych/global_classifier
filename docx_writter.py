from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

document.add_picture('./abc.png', width=Inches(1.25))

document.add_page_break()

document.save('demo.docx')