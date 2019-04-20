from sklearn import metrics
from sklearn.metrics import roc_curve, auc
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import GradientBoostingClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.svm import SVC
from sklearn.ensemble import RandomForestClassifier

from sklearn.model_selection import train_test_split
from sklearn.model_selection import GridSearchCV

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import scikitplot as skplt

import behavior.data_parameter as pr
import data_loader as dl
import full_method as fm

path = './initial_data/Cow_man.csv'
ActualData = pd.read_csv(path, usecols=['time', 'gFx', 'gFy', 'gFz', 'Class'], sep=',', low_memory=False)
time = ActualData['time']
Ax = ActualData['gFx']
Ay = ActualData['gFy']
Az = ActualData['gFz']
Class = ActualData['Class']

size_window = 600


def switch_data_parameter(x):
    global res

    if x == 1:
        res = pr.meanpar(Ax, size_window)
    elif x == 2:
        res = pr.meanpar(Ay, size_window)
    elif x == 3:
        res = pr.meanpar(Az, size_window)
    elif x == 4:
        res = pr.sd(Ax, size_window)
    elif x == 5:
        res = pr.sd(Ay, size_window)
    elif x == 6:
        res = pr.sd(Az, size_window)
    elif x == 7:
        res = pr.meanpar(pr.odba(Ax, Ay, Az), size_window)
    elif x == 8:
        res = pr.meanpar(pr.vedba(Ax, Ay, Az), size_window)
    elif x == 9:
        res = pr.quantiles(Ax, size_window, 0.1)
    elif x == 10:
        res = pr.quantiles(Ay, size_window, 0.1)
    elif x == 11:
        res = pr.quantiles(Az, size_window, 0.1)
    elif x == 12:
        res = pr.quantiles(Ax, size_window, 0.25)
    elif x == 13:
        res = pr.quantiles(Ay, size_window, 0.25)
    elif x == 14:
        res = pr.quantiles(Az, size_window, 0.25)
    elif x == 15:
        res = pr.quantiles(Ax, size_window, 0.5)
    elif x == 16:
        res = pr.quantiles(Ay, size_window, 0.5)
    elif x == 17:
        res = pr.quantiles(Az, size_window, 0.5)
    elif x == 18:
        res = pr.quantiles(Ax, size_window, 0.75)
    elif x == 19:
        res = pr.quantiles(Ay, size_window, 0.75)
    elif x == 20:
        res = pr.quantiles(Az, size_window, 0.75)
    elif x == 21:
        res = pr.quantiles(Ax, size_window, 0.9)
    elif x == 22:
        res = pr.quantiles(Ay, size_window, 0.9)
    elif x == 23:
        res = pr.quantiles(Az, size_window, 0.9)
    elif x == 24:
        res = pr.skewness(Ax, size_window)
    elif x == 25:
        res = pr.skewness(Ay, size_window)
    elif x == 26:
        res = pr.skewness(Az, size_window)
    elif x == 27:
        res = pr.kurtosis(Ax, size_window)
    elif x == 28:
        res = pr.kurtosis(Ay, size_window)
    elif x == 29:
        res = pr.kurtosis(Az, size_window)
    elif x == 30:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.1, 1)
    elif x == 31:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.1, 1)
    elif x == 32:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.1, 1)
    elif x == 33:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.1, 2)
    elif x == 34:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.1, 2)
    elif x == 35:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.1, 2)
    elif x == 36:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.1, 3)
    elif x == 37:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.1, 3)
    elif x == 38:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.1, 3)
    elif x == 39:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.1, 4)
    elif x == 40:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.1, 4)
    elif x == 41:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.1, 4)
    elif x == 42:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.5, 1)
    elif x == 43:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.5, 1)
    elif x == 44:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.5, 1)
    elif x == 45:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.5, 2)
    elif x == 46:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.5, 2)
    elif x == 47:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.5, 2)
    elif x == 48:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.5, 3)
    elif x == 49:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.5, 3)
    elif x == 50:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.5, 3)
    elif x == 51:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.5, 4)
    elif x == 52:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.5, 4)
    elif x == 53:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.5, 4)
    elif x == 54:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.9, 1)
    elif x == 55:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.9, 1)
    elif x == 56:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.9, 1)
    elif x == 57:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.9, 2)
    elif x == 58:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.9, 2)
    elif x == 59:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.9, 2)
    elif x == 60:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.9, 3)
    elif x == 61:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.9, 3)
    elif x == 62:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.9, 3)
    elif x == 63:
        res = pr.sumQuantilesMinMax(Ax, size_window, 0.9, 4)
    elif x == 64:
        res = pr.sumQuantilesMinMax(Ay, size_window, 0.9, 4)
    elif x == 65:
        res = pr.sumQuantilesMinMax(Az, size_window, 0.9, 4)
    elif x == 66:
        res = pr.pairwiseDifferences(Ax, size_window, 0.05, 0.1)
    elif x == 67:
        res = pr.pairwiseDifferences(Ay, size_window, 0.05, 0.1)
    elif x == 68:
        res = pr.pairwiseDifferences(Az, size_window, 0.05, 0.1)
    elif x == 69:
        res = pr.pairwiseDifferences(Ax, size_window, 0.1, 0.25)
    elif x == 70:
        res = pr.pairwiseDifferences(Ay, size_window, 0.1, 0.25)
    elif x == 71:
        res = pr.pairwiseDifferences(Az, size_window, 0.1, 0.25)
    elif x == 72:
        res = pr.pairwiseDifferences(Ax, size_window, 0.25, 0.5)
    elif x == 73:
        res = pr.pairwiseDifferences(Ay, size_window, 0.25, 0.5)
    elif x == 74:
        res = pr.pairwiseDifferences(Az, size_window, 0.25, 0.5)
    elif x == 75:
        res = pr.pairwiseDifferences(Ax, size_window, 0.5, 0.75)
    elif x == 76:
        res = pr.pairwiseDifferences(Ay, size_window, 0.5, 0.75)
    elif x == 77:
        res = pr.pairwiseDifferences(Az, size_window, 0.5, 0.75)
    elif x == 78:
        res = pr.pairwiseDifferences(Ax, size_window, 0.75, 0.9)
    elif x == 79:
        res = pr.pairwiseDifferences(Ay, size_window, 0.75, 0.9)
    elif x == 80:
        res = pr.pairwiseDifferences(Az, size_window, 0.75, 0.9)
    elif x == 81:
        res = pr.AAD(Ax, Ay, Az, size_window)
    elif x == 82:
        res = pr.meanpar(pr.averageIntensity(Ax, Ay, Az), size_window)
    elif x == 83:
        res = pr.signalMagnitudeArea(Ax, Ay, Az, size_window)
    elif x == 84:
        res = pr.movementVariation(Ax, Ay, Az, size_window)
    elif x == 85:
        res = pr.entropy(Ax, Ay, Az, size_window)
    elif x == 86:
        res = pr.energy(Ax, Ay, Az, size_window)
    elif x == 87:
        res = pr.activity(Ax, size_window)
    elif x == 88:
        res = pr.activity(Ay, size_window)
    elif x == 89:
        res = pr.activity(Az, size_window)
    elif x == 90:
        res = pr.mobility(Ax, size_window)
    elif x == 91:
        res = pr.mobility(Ay, size_window)
    elif x == 92:
        res = pr.mobility(Az, size_window)
    elif x == 93:
        res = pr.complexity(Ax, size_window)
    elif x == 94:
        res = pr.complexity(Ay, size_window)
    elif x == 95:
        res = pr.complexity(Az, size_window)
    elif x == 96:
        res = pr.dwa(Ax, size_window)
    elif x == 97:
        res = pr.dwa(Ay, size_window)
    elif x == 98:
        res = pr.dwa(Az, size_window)
    elif x == 99:
        res = pr.meanpar(pr.correlation(Ax, Ay, Az, 1), size_window)
    elif x == 100:
        res = pr.meanpar(pr.correlation(Ax, Ay, Az, 2), size_window)
    elif x == 101:
        res = pr.meanpar(pr.correlation(Ax, Ay, Az, 3), size_window)
    elif x == 102:
        res = pr.meanclass(Class, size_window)

    return res


def switch_name_data_parameter(x):
    return {
        1: 'ax_mean',
        2: 'ay_mean',
        3: 'az_mean',
        4: 'standard_deviation_x',
        5: 'standard_deviation_y',
        6: 'standard_deviation_z',
        7: 'overall_dynamic_body',
        8: 'vectorial_dynamic_body',
        9: 'quan_01_x',
        10: 'quan_01_y',
        11: 'quan_01_z',
        12: 'quan_025_x',
        13: 'quan_025_y',
        14: 'quan_025_z',
        15: 'quan_05_x',
        16: 'quan_05_y',
        17: 'quan_05_z',
        18: 'quan_075_x',
        19: 'quan_075_y',
        20: 'quan_075_z',
        21: 'quan_09_x',
        22: 'quan_09_y',
        23: 'quan_09_z',
        24: 'skewness_x',
        25: 'skewness_y',
        26: 'skewness_z',
        27: 'kurtosis_x',
        28: 'kurtosis_y',
        29: 'kurtosis_z',
        30: 'sum_quantiles_min_01_x',
        31: 'sum_quantiles_min_01_y',
        32: 'sum_quantiles_min_01_z',
        33: 'sum_quantiles_min_squares_01_x',
        34: 'sum_quantiles_min_squares_01_y',
        35: 'sum_quantiles_min_squares_01_z',
        36: 'sum_quantiles_max_01_x',
        37: 'sum_quantiles_max_01_y',
        38: 'sum_quantiles_max_01_z',
        39: 'sum_quantiles_max_squares_01_x',
        40: 'sum_quantiles_max_squares_01_y',
        41: 'sum_quantiles_max_squares_01_z',
        42: 'sum_quantiles_min_05_x',
        43: 'sum_quantiles_min_05_y',
        44: 'sum_quantiles_min_05_z',
        45: 'sum_quantiles_min_squares_05_x',
        46: 'sum_quantiles_min_squares_05_y',
        47: 'sum_quantiles_min_squares_05_z',
        48: 'sum_quantiles_max_05_x',
        49: 'sum_quantiles_max_05_y',
        50: 'sum_quantiles_max_05_z',
        51: 'sum_quantiles_max_squares_05_x',
        52: 'sum_quantiles_max_squares_05_y',
        53: 'sum_quantiles_max_squares_05_z',
        54: 'sum_quantiles_min_09_x',
        55: 'sum_quantiles_min_09_y',
        56: 'sum_quantiles_min_09_z',
        57: 'sum_quantiles_min_squares_09_x',
        58: 'sum_quantiles_min_squares_09_y',
        59: 'sum_quantiles_min_squares_09_z',
        60: 'sum_quantiles_max_09_x',
        61: 'sum_quantiles_max_09_y',
        62: 'sum_quantiles_max_09_z',
        63: 'sum_quantiles_max_squares_09_x',
        64: 'sum_quantiles_max_squares_09_y',
        65: 'sum_quantiles_max_squares_09_z',
        66: 'pairwise_differences_005_01_x',
        67: 'pairwise_differences_005_01_y',
        68: 'pairwise_differences_005_01_z',
        69: 'pairwise_differences_01_025_x',
        70: 'pairwise_differences_01_025_y',
        71: 'pairwise_differences_01_025_z',
        72: 'pairwise_differences_025_05_x',
        73: 'pairwise_differences_025_05_y',
        74: 'pairwise_differences_025_05_z',
        75: 'pairwise_differences_05_075_x',
        76: 'pairwise_differences_05_075_y',
        77: 'pairwise_differences_05_075_z',
        78: 'pairwise_differences_075_09_x',
        79: 'pairwise_differences_075_09_y',
        80: 'pairwise_differences_075_09_z',
        81: 'average_absolute_difference',
        82: 'average_intensity',
        83: 'signal_magnitude_area',
        84: 'movement_variation',
        85: 'entropy',
        86: 'energy',
        87: 'activity_x',
        88: 'activity_y',
        89: 'activity_z',
        90: 'mobility_x',
        91: 'mobility_y',
        92: 'mobility_z',
        93: 'complexity_x',
        94: 'complexity_y',
        95: 'complexity_z',
        96: 'durbin_watson_x',
        97: 'durbin_watson_y',
        98: 'durbin_watson_z',
        99: 'correlation_magnitude_horizontal',
        100: 'correlation_magnitude_vertical',
        101: 'correlation_horizontal_vertical',
        102: 'Activity'
    }[x]


def generate():
    document = Document()

    for i in range(101):
        for j in range(101):
            d = {switch_name_data_parameter(i + 1): switch_data_parameter(i + 1),
                 switch_name_data_parameter(j + 1): switch_data_parameter(j + 1),
                 switch_name_data_parameter(102): switch_data_parameter(102)}

            full = pd.DataFrame(data=d)
            train, test = train_test_split(full, test_size=0.50)

            train.to_csv('train.csv', index=False)
            test.to_csv('test.csv', index=False)

            train = dl.data_loading('./train.csv')
            test = dl.data_loading('./test.csv')
            labels = dl.labels_loading()

            # get X_train and y_train from csv files
            X_train = train.drop(['Activity'], axis=1)
            y_train = train.Activity

            # get X_test and y_test from test csv file
            X_test = test.drop(['Activity'], axis=1)
            y_test = test.Activity

            document.add_heading(
                'Parameters: ' + switch_name_data_parameter(i + 1) + ' and ' + switch_name_data_parameter(j + 1), 0)

            # DT classifier
            parameters_dt = {'max_depth': np.arange(3, 10, 2)}
            dt = DecisionTreeClassifier()
            dt_grid = GridSearchCV(dt, param_grid=parameters_dt, n_jobs=-1)
            dt_grid_results = fm.perform_model(dt_grid, X_train, y_train, X_test, y_test, class_labels=labels)

            document.add_paragraph(
                format('Сlassification algorithm: ' + 'Decision Trees' + '\n' + 'Full accuracy: ' + str(
                    dt_grid_results['accuracy'])), style='List Number'
            )

            document.add_picture('./abc.png', width=Inches(3.25))

            # Gradient Boost DT classifier
            param_grid = {'max_depth': np.arange(5, 8, 1), \
                          'n_estimators': np.arange(130, 170, 10)}
            gbdt = GradientBoostingClassifier()
            gbdt_grid = GridSearchCV(gbdt, param_grid=param_grid, n_jobs=-1)
            gbdt_grid_results = fm.perform_model(gbdt_grid, X_train, y_train, X_test, y_test, class_labels=labels)

            document.add_paragraph(
                format('Сlassification algorithm: ' + 'Decision Trees' + '\n' + 'Full accuracy: ' + str(
                    gbdt_grid_results['accuracy'])), style='List Number'
            )

            document.add_picture('./abc.png', width=Inches(3.25))

            # Logistic Regression classifier
            parameters_lr = {'C': [0.001, 0.1, 1, 10, 20, 30], 'penalty': ['l2', 'l1']}
            log_reg = LogisticRegression(multi_class='auto')
            log_reg_grid = GridSearchCV(log_reg, param_grid=parameters_lr, cv=3, verbose=1, n_jobs=-1)
            log_reg_grid_results = fm.perform_model(log_reg_grid, X_train, y_train, X_test, y_test, class_labels=labels)

            document.add_paragraph(
                format('Сlassification algorithm: ' + 'Decision Trees' + '\n' + 'Full accuracy: ' + str(
                    log_reg_grid_results['accuracy'])), style='List Number'
            )

            # Random Forest classifier
            parameters_rf = {'n_estimators': np.arange(10, 201, 20), 'max_depth': np.arange(3, 15, 2)}
            rfc = RandomForestClassifier()
            rfc_grid = GridSearchCV(rfc, param_grid=parameters_rf, n_jobs=-1)
            rfc_grid_results = fm.perform_model(rfc_grid, X_train, y_train, X_test, y_test, class_labels=labels)

            document.add_paragraph(
                format('Сlassification algorithm: ' + 'Decision Trees' + '\n' + 'Full accuracy: ' + str(
                    rfc_grid_results['accuracy'])), style='List Number'
            )

            document.add_picture('./abc.png', width=Inches(3.25))

            document.add_page_break()

    document.save('demo.docx')


generate()

print('done')
